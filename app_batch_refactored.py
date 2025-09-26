"""
=============================================================================
TOOL BATCH - ĐIỀN NHIỀU GIẤY XÁC NHẬN TÌNH TRẠNG HÔN NHÂN
=============================================================================
Ứng dụng Streamlit để xử lý hàng loạt giấy xác nhận tình trạng hôn nhân
Author: AI Assistant
Version: 2.0 (Refactored)
=============================================================================
"""

import streamlit as st

# Cấu hình Streamlit - PHẢI ĐẶT ĐẦU TIÊN
st.set_page_config(
    page_title="Tool Điền Giấy Xác Nhận Batch",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="collapsed"
)

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import tempfile
import os
import zipfile
from io import BytesIO
import uuid
import time
import atexit

# =============================================================================
# CONSTANTS & CONFIGURATION
# =============================================================================

TEMP_DIR = tempfile.gettempdir()
MAX_FILES = 5
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

# Session management
SESSION_ID = str(uuid.uuid4())[:8]
TIMESTAMP = int(time.time())
SESSION_FILES = []

def cleanup_session_files():
    """Cleanup temporary files for current session"""
    for file_path in SESSION_FILES:
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except:
            pass

# Register cleanup on exit
atexit.register(cleanup_session_files)

def get_unique_temp_path(prefix, extension=".docx"):
    """Generate unique temporary file path"""
    filename = f"{prefix}_{SESSION_ID}_{TIMESTAMP}_{int(time.time() * 1000000) % 1000000}{extension}"
    path = os.path.join(TEMP_DIR, filename)
    SESSION_FILES.append(path)
    return path

# Danh sách từ khóa cần loại bỏ khi tìm tên người ký
BLACKLIST_KEYWORDS = [
    'CHỦ TỊCH', 'PHÓ CHỦ TỊCH', 'KT.', 'GIẤY', 'XÁC NHẬN', 'TÌNH TRẠNG', 
    'HÔN NHÂN', 'UBND', 'ỦY BAN', 'NHÂN DÂN', 'SỞ', 'PHÒNG', 'BAN',
    'CỘNG HÒA', 'XÃ HỘI', 'CHỦ NGHĨA', 'VIỆT NAM', 'ĐỘC LẬP', 'TỰ DO',
    'HẠNH PHÚC', 'TỈNH', 'THÀNH PHỐ', 'QUẬN', 'HUYỆN', 'XÃ', 'PHƯỜNG'
]

# Họ phổ biến Việt Nam
COMMON_SURNAMES = [
    'Nguyễn', 'Trần', 'Lê', 'Phạm', 'Hoàng', 'Huỳnh', 'Phan', 'Vũ', 
    'Võ', 'Đặng', 'Bùi', 'Đỗ', 'Hồ', 'Ngô', 'Dương'
]

# =============================================================================
# UTILITY FUNCTIONS
# =============================================================================

def is_vietnamese_name(text):
    """
    Kiểm tra xem text có phải tên người Việt Nam không
    
    Args:
        text (str): Chuỗi cần kiểm tra
        
    Returns:
        bool: True nếu là tên người Việt Nam hợp lệ
    """
    if not text or len(text.strip()) < 3:
        return False
    
    text = text.strip()
    
    # Loại bỏ các từ khóa công văn
    for word in BLACKLIST_KEYWORDS:
        if word in text.upper():
            return False
    
    # Kiểm tra pattern tên Việt Nam (2-5 từ, mỗi từ bắt đầu bằng chữ hoa)
    words = text.split()
    if len(words) < 2 or len(words) > 5:
        return False
    
    for word in words:
        if not re.match(r'^[A-ZÀ-Ỹ][a-zà-ỹ]*$', word):
            return False
    
    # Không chứa số hoặc ký tự đặc biệt
    if re.search(r'[\d\.\,\:\;\!\?\(\)\[\]\{\}]', text):
        return False
    
    return True

def score_name_candidate(name, context, all_lines):
    """
    Chấm điểm ứng viên tên để chọn tên tốt nhất
    
    Args:
        name (str): Tên ứng viên
        context (str): Dòng chứa tên
        all_lines (list): Tất cả các dòng trong văn bản
        
    Returns:
        int: Điểm số của ứng viên
    """
    score = 10  # Điểm cơ bản
    
    # Ưu tiên tên ở cuối văn bản
    try:
        line_index = all_lines.index(context)
        total_lines = len(all_lines)
        if line_index >= total_lines - 3:
            score += 20
        elif line_index >= total_lines - 5:
            score += 10
    except:
        pass
    
    # Ưu tiên tên sau chức vụ
    if re.search(r'(CHỦ TỊCH|PHÓ CHỦ TỊCH|KT\.)', context, re.IGNORECASE):
        score += 15
    
    # Ưu tiên tên có độ dài phù hợp
    word_count = len(name.split())
    if word_count == 3:
        score += 15
    elif word_count == 2:
        score += 10
    elif word_count == 4:
        score += 5
    
    # Trừ điểm nếu tên quá ngắn hoặc quá dài
    if len(name) < 6:
        score -= 5
    elif len(name) > 25:
        score -= 10
    
    # Ưu tiên họ phổ biến Việt Nam
    first_word = name.split()[0]
    if first_word in COMMON_SURNAMES:
        score += 10
    
    return score

def validate_file(file_path):
    """
    Kiểm tra tính hợp lệ của file
    
    Args:
        file_path (str): Đường dẫn file
        
    Returns:
        tuple: (is_valid, error_message)
    """
    if not os.path.exists(file_path):
        return False, "File không tồn tại"
        
    file_size = os.path.getsize(file_path)
    if file_size == 0:
        return False, "File rỗng"
    if file_size > MAX_FILE_SIZE:
        return False, "File quá lớn (>50MB)"
        
    try:
        doc = Document(file_path)
        return True, None
    except Exception as e:
        return False, f"File không hợp lệ: {str(e)}"

# =============================================================================
# DATA EXTRACTION FUNCTIONS
# =============================================================================

def extract_text_from_document(doc_path):
    """
    Trích xuất text từ file Word
    
    Args:
        doc_path (str): Đường dẫn file Word
        
    Returns:
        str: Nội dung text của file
    """
    doc = Document(doc_path)
    
    # Lấy text từ paragraphs
    full_text = '\n'.join([para.text for para in doc.paragraphs])
    
    # Lấy text từ tables
    table_text = ''
    try:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + '\n'
    except Exception:
        pass
    
    return full_text + '\n' + table_text

def find_person_signature(all_text):
    """
    Tìm tên người ký bằng thuật toán nâng cao
    
    Args:
        all_text (str): Toàn bộ nội dung văn bản
        
    Returns:
        tuple: (ten_nguoi_ky, chuc_vu)
    """
    # Tìm chức vụ
    chuc_vu = ''
    if re.search(r'KT\.\s*CHỦ TỊCH\s*PHÓ CHỦ TỊCH', all_text):
        chuc_vu = 'KT. CHỦ TỊCH - PHÓ CHỦ TỊCH'
    elif re.search(r'PHÓ CHỦ TỊCH', all_text):
        chuc_vu = 'PHÓ CHỦ TỊCH'
    elif re.search(r'CHỦ TỊCH', all_text):
        chuc_vu = 'CHỦ TỊCH'
    
    # Thuật toán tìm tên nâng cao
    ten_nguoi_ky = ''
    
    # Bước 1: Tách văn bản thành các dòng
    lines = [line.strip() for line in all_text.split('\n') if line.strip()]
    
    # Bước 2: Tìm vị trí chức vụ cuối cùng
    chuc_vu_positions = []
    for i, line in enumerate(lines):
        if re.search(r'(KT\.|CHỦ TỊCH|PHÓ CHỦ TỊCH)', line, re.IGNORECASE):
            chuc_vu_positions.append(i)
    
    # Bước 3: Tìm tên sau vị trí chức vụ cuối cùng
    if chuc_vu_positions:
        start_search = chuc_vu_positions[-1] + 1
        
        for i in range(start_search, min(start_search + 5, len(lines))):
            if i < len(lines):
                candidate = lines[i].strip()
                if is_vietnamese_name(candidate):
                    ten_nguoi_ky = candidate
                    break
    
    # Bước 4: Tìm trong toàn bộ văn bản nếu chưa có
    if not ten_nguoi_ky:
        name_candidates = []
        
        for line in lines:
            words = line.split()
            for i in range(len(words)):
                for j in range(i+2, min(i+6, len(words)+1)):
                    candidate = ' '.join(words[i:j])
                    if is_vietnamese_name(candidate):
                        name_candidates.append((candidate, line))
        
        if name_candidates:
            scored_candidates = []
            for name, context in name_candidates:
                score = score_name_candidate(name, context, lines)
                scored_candidates.append((score, name))
            
            scored_candidates.sort(reverse=True)
            ten_nguoi_ky = scored_candidates[0][1]
    
    return ten_nguoi_ky, chuc_vu

def extract_field_data(all_text, field_patterns):
    """
    Trích xuất dữ liệu các trường theo patterns
    
    Args:
        all_text (str): Nội dung văn bản
        field_patterns (dict): Dictionary chứa patterns cho từng trường
        
    Returns:
        dict: Dữ liệu đã trích xuất
    """
    data = {}
    
    for field_name, patterns in field_patterns.items():
        data[field_name] = ''
        
        if isinstance(patterns, list):
            for pattern in patterns:
                match = re.search(pattern, all_text)
                if match:
                    data[field_name] = match.group(1).strip()
                    break
        else:
            match = re.search(patterns, all_text)
            if match:
                data[field_name] = match.group(1).strip()
    
    return data

def sanitize_filename(filename):
    """Làm sạch tên file để tránh lỗi"""
    import string
    # Loại bỏ ký tự đặc biệt, chỉ giữ chữ, số, dấu gạch ngang, gạch dưới, chấm
    valid_chars = "-_. %s%s" % (string.ascii_letters, string.digits)
    filename = ''.join(c for c in filename if c in valid_chars)
    # Thay thế khoảng cách bằng gạch dưới
    filename = re.sub(r'\s+', '_', filename)
    return filename

def extract_data_from_input(input_path):
    """
    Trích xuất dữ liệu từ file input
    
    Args:
        input_path (str): Đường dẫn file input
        
    Returns:
        tuple: (data_dict, error_message)
    """
    try:
        # Validate file
        is_valid, error = validate_file(input_path)
        if not is_valid:
            return None, error
        
        # Extract text
        all_text = extract_text_from_document(input_path)
        
        if not all_text.strip():
            return None, "File không có nội dung"
        
        # Kiểm tra loại file
        if not re.search(r'GIẤY XÁC NHẬN TÌNH TRẠNG HÔN NHÂN', all_text, re.IGNORECASE):
            return None, "File không phải Giấy xác nhận tình trạng hôn nhân"
        
        # Define field patterns
        field_patterns = {
            'Số': [r'Số:\s*([\w/\-]+)', r'Số\s*:\s*([\w/\-]+)'],
            'Họ tên': r'Họ, chữ đệm, tên:\s*([A-ZÀ-Ỹ\s]+?)(?=\s*Ngày|$)',
            'Ngày sinh': r'Ngày, tháng, năm sinh:\s*(\d+/\d+/\d+)',
            'Giới tính': r'Giới tính:\s*([^\n\r]+?)(?=\s*(?:Dân tộc|$))',
            'Dân tộc': r'Dân tộc:\s*([^\n\r]+?)(?=\s*(?:Quốc tịch|$))',
            'Quốc tịch': r'Quốc tịch:\s*([^\n\r]+?)(?=\s*(?:Giấy|Nơi|$))',
            'Nơi cư trú': r'Nơi cư trú:\s*(.+?)(?=\s*Tình trạng|$)',
            'Giấy tờ tùy thân': r'Giấy tờ tùy thân:\s*(.+?)(?=\s*Nơi|$)',
            'Tình trạng hôn nhân': r'Tình trạng hôn nhân:\s*(.+?)(?=\s*Giấy|$)',
            'Mục đích sử dụng': r'sử dụng để:\s*(.+?)(?=\s*Giấy|$)'
        }
        
        # Extract basic fields
        data = extract_field_data(all_text, field_patterns)
        
        # Extract date
        try:
            date_match = re.search(r'ngày\s*(\d+)\s*tháng\s*(\d+)\s*năm\s*(\d+)', all_text)
            data['Ngày cấp'] = f"{date_match.group(1)}/{date_match.group(2)}/{date_match.group(3)}" if date_match else ''
        except:
            data['Ngày cấp'] = ''
        
        # Extract person signature
        ten_nguoi_ky, chuc_vu = find_person_signature(all_text)
        if ten_nguoi_ky and chuc_vu:
            data['Người ký'] = f"{ten_nguoi_ky} - {chuc_vu}"
        elif ten_nguoi_ky:
            data['Người ký'] = ten_nguoi_ky
        elif chuc_vu:
            data['Người ký'] = chuc_vu
        else:
            data['Người ký'] = ''
        
        # Set người đề nghị
        data['Người đề nghị'] = data['Họ tên']
        
        # Clean data
        for key in data:
            if isinstance(data[key], str):
                data[key] = data[key].strip()
        
        # Check required fields
        required_fields = ['Số', 'Ngày cấp', 'Họ tên', 'Ngày sinh', 'Giới tính', 
                          'Dân tộc', 'Quốc tịch', 'Nơi cư trú', 'Giấy tờ tùy thân', 
                          'Tình trạng hôn nhân', 'Mục đích sử dụng', 'Người ký']
        
        missing_fields = [field for field in required_fields if not data.get(field)]
        
        if missing_fields:
            error_msg = f"Thiếu dữ liệu bắt buộc: {', '.join(missing_fields)}"
            return data, error_msg
        
        return data, None
        
    except Exception as e:
        return None, f"Lỗi không xác định: {str(e)}"
# =============================================================================
# TEMPLATE FILLING FUNCTIONS
# =============================================================================

def fill_template(template_path, data, output_docx_path):
    """
    Điền dữ liệu vào template
    
    Args:
        template_path (str): Đường dẫn file template
        data (dict): Dữ liệu cần điền
        output_docx_path (str): Đường dẫn file output
        
    Returns:
        bool: True nếu thành công
    """
    try:
        if not os.path.exists(template_path):
            return False
            
        doc = Document(template_path)
        
        # Fill data in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    try:
                        cell_text = cell.text
                        
                        # Replace specific patterns
                        if 'Số:' in cell_text and data.get('Số'):
                            cell.text = re.sub(r'Số:\s*[.………_\-]+', f"Số: {data['Số']}", cell_text, count=1)
                        
                        if 'Ngày, tháng, năm cấp:' in cell_text and data.get('Ngày cấp'):
                            cell.text = re.sub(r'Ngày, tháng, năm cấp:\s*[.………/\-]+', f"Ngày, tháng, năm cấp: {data['Ngày cấp']}", cell_text, count=1)
                        
                        if 'Họ, chữ đệm, tên:' in cell_text and data.get('Họ tên'):
                            cell.text = re.sub(r'Họ, chữ đệm, tên:\s*[.…………]+', f"Họ, chữ đệm, tên: {data['Họ tên']}", cell_text)
                        
                        if 'Họ, chữ đệm, tên, chức vụ người ký' in cell_text and data.get('Người ký'):
                            cell.text = re.sub(r'Họ, chữ đệm, tên, chức vụ người ký[^:]*:\s*[.…………]+', f"Họ, chữ đệm, tên, chức vụ người ký Giấy xác nhận tình trạng hôn nhân: {data['Người ký']}", cell_text, count=1)
                        
                        # Flexible string replacement for different dot formats
                        if 'Giới tính:' in cell_text and data.get('Giới tính'):
                            # Try multiple dot patterns
                            patterns = ['Giới tính: …………….', 'Giới tính:…………….']
                            for pattern in patterns:
                                if pattern in cell_text:
                                    cell.text = cell_text.replace(pattern, f"Giới tính: {data['Giới tính']}")
                                    break
                        
                        if 'Dân tộc:' in cell.text and data.get('Dân tộc'):
                            patterns = ['Dân tộc: …………….', 'Dân tộc:…………….']
                            for pattern in patterns:
                                if pattern in cell.text:
                                    cell.text = cell.text.replace(pattern, f"Dân tộc: {data['Dân tộc']}")
                                    break
                        
                        if 'Quốc tịch:' in cell.text and data.get('Quốc tịch'):
                            patterns = ['Quốc tịch: …………….', 'Quốc tịch:…………….']
                            for pattern in patterns:
                                if pattern in cell.text:
                                    cell.text = cell.text.replace(pattern, f"Quốc tịch: {data['Quốc tịch']}")
                                    break
                        
                        # Fill other fields
                        field_mappings = [
                            ('Ngày, tháng, năm sinh:', 'Ngày sinh'),
                            ('Nơi cưu trú:', 'Nơi cư trú'),
                            ('Giấy tờ tùy thân:', 'Giấy tờ tùy thân'),
                            ('Tình trạng hôn nhân:', 'Tình trạng hôn nhân'),
                            ('Mục đích sử dụng:', 'Mục đích sử dụng')
                        ]
                        
                        for field_name, data_key in field_mappings:
                            if field_name in cell_text and data.get(data_key):
                                # Simple pattern like original code
                                pattern = field_name.replace(':', r':\s*[.…………]+')
                                cell.text = re.sub(pattern, f"{field_name} {data[data_key]}", cell_text, count=1)
                    except:
                        continue
        
        # Set font formatting
        try:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # Căn phải cho ngày cấp
                            if 'Ngày, tháng, năm cấp:' in cell.text:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            
                            for run in paragraph.runs:
                                try:
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(13)
                                    if 'Họ, chữ đệm, tên, chức vụ người ký' in cell.text:
                                        run.font.bold = True
                                except:
                                    continue
        except:
            pass
        
        doc.save(output_docx_path)
        return True
        
    except Exception:
        return False

# =============================================================================
# STREAMLIT UI FUNCTIONS
# =============================================================================

def render_custom_css():
    """Render custom CSS cho giao diện"""
    st.markdown("""
    <style>
        .main-header {
            background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
            padding: 2rem;
            border-radius: 10px;
            margin-bottom: 2rem;
            text-align: center;
            color: white;
        }
        .upload-section {
            background: #f8f9fa;
            padding: 1.5rem;
            border-radius: 10px;
            border-left: 4px solid #007bff;
            margin: 1rem 0;
        }
        .success-box {
            background: #d4edda;
            border: 1px solid #c3e6cb;
            border-radius: 8px;
            padding: 1rem;
            margin: 0.5rem 0;
        }
        .error-box {
            background: #f8d7da;
            border: 1px solid #f5c6cb;
            border-radius: 8px;
            padding: 1rem;
            margin: 0.5rem 0;
        }
        .info-box {
            background: #e2f3ff;
            border: 1px solid #b8daff;
            border-radius: 8px;
            padding: 1rem;
            margin: 0.5rem 0;
        }
        .stats-card {
            background: white;
            padding: 1rem;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            text-align: center;
            margin: 0.5rem;
        }
    </style>
    """, unsafe_allow_html=True)

def render_header():
    """Render header chính"""
    st.markdown(f"""
    <div class="main-header">
        <h1>🏛️ Tool Batch - Xử Lý Giấy Xác Nhận</h1>
        <p>Điền nhiều giấy xác nhận tình trạng hôn nhân cùng lúc một cách nhanh chóng và chính xác</p>
        <small style="opacity: 0.7;">Session: {SESSION_ID}</small>
    </div>
    """, unsafe_allow_html=True)

def render_file_upload_section():
    """Render section upload file dữ liệu"""
    st.markdown("""
    <div class="upload-section">
        <h3> Bước 1: Upload File Dữ Liệu</h3>
        <p>Chọn tối đa 5 file .docx chứa thông tin cần điền</p>
    </div>
    """, unsafe_allow_html=True)
    
    return st.file_uploader(
        "", 
        type="docx", 
        accept_multiple_files=True,
        help="Kéo thả hoặc click để chọn file (tối đa 5MB mỗi file)",
        key="input_files"
    )

def render_template_upload_section():
    """Hiển thị thông tin template cố định"""
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown("""
    <div class="upload-section">
        <h3>📋 Template Cố Định</h3>
        <p> Tool sử dụng template được tối ưu hóa cho giấy xác nhận tình trạng hôn nhân</p>
        <p> Không cần upload template - đã được cài đặt sẵn</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Trả về đường dẫn template cố định
    return "temp/mau.docx"

def display_file_stats(valid_count, error_count):
    """Hiển thị thống kê file"""
    col1, col2 = st.columns(2)
    with col1:
        st.markdown(f"""
        <div class="stats-card" style="background: #d4edda;">
            <h3 style="color: #155724;">{valid_count}</h3>
            <p>File hợp lệ</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown(f"""
        <div class="stats-card" style="background: #f8d7da;">
            <h3 style="color: #721c24;">{error_count}</h3>
            <p>File có lỗi</p>
        </div>
        """, unsafe_allow_html=True)

def display_data_details(data_list, error_list):
    """Hiển thị chi tiết dữ liệu"""
    if data_list:
        with st.expander(f" Xem chi tiết {len(data_list)} file hợp lệ", expanded=False):
            for i, data in enumerate(data_list):
                st.markdown(f"**📄 {data['file_name']}**")
                
                col1, col2 = st.columns(2)
                data_items = list(data.items())
                mid = len(data_items) // 2
                
                with col1:
                    for k, v in data_items[:mid]:
                        if k not in ['file_name', 'file_index']:
                            st.write(f"**{k}:** {v}")
                
                with col2:
                    for k, v in data_items[mid:]:
                        if k not in ['file_name', 'file_index']:
                            st.write(f"**{k}:** {v}")
                
                if i < len(data_list) - 1:
                    st.divider()
    
    if error_list:
        with st.expander(f"❌ Xem chi tiết {len(error_list)} file có lỗi", expanded=False):
            for error_info in error_list:
                st.markdown(f"""
                <div class="error-box">
                    <h4>📄 {error_info['file_name']}</h4>
                    <p><strong>Lỗi:</strong> {error_info['error']}</p>
                </div>
                """, unsafe_allow_html=True)
                
                if error_info['data']:
                    st.write("**Dữ liệu đọc được (không đầy đủ):**")
                    col1, col2 = st.columns(2)
                    data_items = list(error_info['data'].items())
                    mid = len(data_items) // 2
                    
                    with col1:
                        for k, v in data_items[:mid]:
                            if k not in ['file_name', 'file_index']:
                                st.write(f"**{k}:** {v}")
                    
                    with col2:
                        for k, v in data_items[mid:]:
                            if k not in ['file_name', 'file_index']:
                                st.write(f"**{k}:** {v}")

def render_footer():
    """Render footer với hướng dẫn"""
    st.markdown("<br><br>", unsafe_allow_html=True)
    st.markdown("""
    <div style="background: #f8f9fa; padding: 2rem; border-radius: 10px; border-top: 3px solid #007bff;">
        <h3 style="color: #007bff; margin-bottom: 1rem;">💡 Hướng Dẫn Sử Dụng</h3>
        <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(250px, 1fr)); gap: 1rem;">
            <div style="background: white; padding: 1rem; border-radius: 8px; border-left: 4px solid #28a745;">
                <h4 style="color: #28a745; margin: 0;">Bước 1</h4>
                <p style="margin: 0.5rem 0 0 0;">Upload tối đa 5 file dữ liệu (.docx)</p>
            </div>
            <div style="background: white; padding: 1rem; border-radius: 8px; border-left: 4px solid #ffc107;">
                <h4 style="color: #ffc107; margin: 0;">Bước 2</h4>
                <p style="margin: 0.5rem 0 0 0;">Upload 1 file template dùng chung</p>
            </div>
            <div style="background: white; padding: 1rem; border-radius: 8px; border-left: 4px solid #dc3545;">
                <h4 style="color: #dc3545; margin: 0;">Bước 3</h4>
                <p style="margin: 0.5rem 0 0 0;">Nhấn 'Xử Lý' và tải file ZIP</p>
            </div>
        </div>
        <div style="margin-top: 1rem; padding: 1rem; background: #e9ecef; border-radius: 8px;">
            <p style="margin: 0; color: #6c757d; text-align: center;">
                <strong>Lưu ý:</strong> Chỉ file hợp lệ mới được xử lý. File có lỗi sẽ được hiển thị chi tiết để bạn có thể sửa chữa.
            </p>
        </div>
    </div>
    """, unsafe_allow_html=True)
# =============================================================================
# MAIN APPLICATION
# =============================================================================

def main():
    """Hàm chính của ứng dụng"""
    
    # Render UI components
    render_custom_css()
    render_header()
    
    # Step 1: Upload input files
    uploaded_inputs = render_file_upload_section()
    
    data_list = []
    error_list = []
    
    if uploaded_inputs:
        # Validate file count
        if len(uploaded_inputs) > MAX_FILES:
            st.error(f"❌ Chỉ được upload tối đa {MAX_FILES} file!")
            uploaded_inputs = uploaded_inputs[:MAX_FILES]
        
        # Process files with progress bar
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for i, uploaded_file in enumerate(uploaded_inputs):
            progress_bar.progress((i + 1) / len(uploaded_inputs))
            status_text.text(f'Đang xử lý: {uploaded_file.name}')
            
            if uploaded_file.name.lower().endswith('.docx'):
                # Tạo tên file tạm thời unique
                input_path = get_unique_temp_path(f"input_{i}")
                try:
                    with open(input_path, "wb") as f:
                        f.write(uploaded_file.getvalue())
                    
                    data, error = extract_data_from_input(input_path)
                    if data and not error:
                        data['file_name'] = uploaded_file.name
                        data['file_index'] = i + 1
                        data_list.append(data)
                    else:
                        error_info = {
                            'file_name': uploaded_file.name,
                            'error': error or "Không đọc được dữ liệu",
                            'data': data
                        }
                        error_list.append(error_info)
                except Exception as e:
                    error_info = {
                        'file_name': uploaded_file.name,
                        'error': f"Lỗi xử lý: {str(e)}",
                        'data': None
                    }
                    error_list.append(error_info)
            else:
                error_info = {
                    'file_name': uploaded_file.name,
                    'error': "Không phải file .docx",
                    'data': None
                }
                error_list.append(error_info)
        
        progress_bar.empty()
        status_text.empty()
        
        # Display results
        display_file_stats(len(data_list), len(error_list))
        display_data_details(data_list, error_list)
    
    # Step 2: Upload template
    uploaded_template = render_template_upload_section()
    template_path = None
    
    if uploaded_template:
        # uploaded_template bây giờ là đường dẫn string, không phải file object
        template_path = uploaded_template
        try:
            # Kiểm tra file template có tồn tại không
            if os.path.exists(template_path):
                test_doc = Document(template_path)
                st.markdown("""
                <div class="success-box">
                    <h4>✅ Template đã sẵn sàng</h4>
                    <p>Template cố định đã được tải thành công</p>
                </div>
                """, unsafe_allow_html=True)
            else:
                raise FileNotFoundError(f"Template file không tồn tại: {template_path}")
        except Exception as e:
            st.markdown(f"""
            <div class="error-box">
                <h4>❌ Template không hợp lệ</h4>
                <p>{str(e)}</p>
            </div>
            """, unsafe_allow_html=True)
            template_path = None
    
    # Step 3: Process files
    st.markdown("<br>", unsafe_allow_html=True)
    
    if data_list and template_path:
        st.markdown("""
        <div class="info-box">
            <h3> Bước 3: Xử Lý File</h3>
            <p>Tất cả đã sẵn sàng! Nhấn nút bên dưới để bắt đầu xử lý</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Display metrics
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("File hợp lệ", len(data_list))
        with col2:
            st.metric("File có lỗi", len(error_list))
        with col3:
            st.metric("Tổng cộng", len(data_list) + len(error_list))
        
        if error_list:
            st.warning(f"⚠️ {len(error_list)} file có lỗi sẽ bị bỏ qua")
        
        # Process button
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            process_button = st.button(
                f" Xử Lý {len(data_list)} File Hợp Lệ", 
                type="primary",
                use_container_width=True
            )
        
        if process_button:
            # Process files with progress tracking
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            zip_buffer = BytesIO()
            used_names = {}
            processed_files = []  # Tạo list để lưu file đã xử lý
            
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                success_count = 0
                
                for i, data in enumerate(data_list):
                    progress_bar.progress((i + 1) / len(data_list))
                    status_text.text(f'Đang xử lý: {data["file_name"]}')
                    
                    try:
                        output_path = get_unique_temp_path(f"output_{i}")
                        
                        if fill_template(template_path, data, output_path):
                            # Generate unique filename với sanitize
                            ho_ten = data.get('Họ tên', f'File_{data["file_index"]}')
                            ho_ten_clean = sanitize_filename(ho_ten)
                            base_name = f"{ho_ten_clean}_GiayXacNhan"
                            
                            if base_name in used_names:
                                used_names[base_name] += 1
                                zip_filename = f"{base_name}_{used_names[base_name]}.docx"
                            else:
                                used_names[base_name] = 1
                                zip_filename = f"{base_name}.docx"
                            
                            # Sanitize zip filename
                            zip_filename = sanitize_filename(zip_filename)
                            
                            with open(output_path, 'rb') as f:
                                zip_file.writestr(zip_filename, f.read())
                            
                            # Thêm vào processed_files
                            processed_files.append((zip_filename, output_path))
                            success_count += 1
                        else:
                            st.error(f"❌ {data['file_name']}: Lỗi khi xử lý template")
                    except Exception as e:
                        st.error(f"❌ {data['file_name']}: {str(e)}")
            
            progress_bar.empty()
            status_text.empty()
            zip_buffer.seek(0)
            
            if success_count > 0:
                st.markdown(f"""
                <div class="success-box">
                    <h3>🎉 Xử Lý Hoàn Thành!</h3>
                    <p>Đã xử lý thành công <strong>{success_count}/{len(data_list)}</strong> file hợp lệ</p>
                </div>
                """, unsafe_allow_html=True)
                
                # Chỉ lấy file thành công (loại bỏ file lỗi)
                success_files = [(name, path) for name, path in processed_files if os.path.exists(path)]
                
                if success_files:
                    # Nút xuất tất cả (chỉ file thành công)
                    st.subheader("📦 Xuất Tất Cả")
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.download_button(
                            f"📄 Tải Tất Cả DOCX ({len(success_files)} file)",
                            zip_buffer.getvalue(),
                            file_name="GiayXacNhan_DOCX.zip",
                            mime="application/zip",
                            use_container_width=True
                        )
                
                    
                    # Tải từng file riêng lẻ (chỉ file thành công)
                    st.subheader("📄 Tải Từng File")
                    for i, (filename, file_path) in enumerate(success_files):
                        if os.path.exists(file_path):  # Double check file tồn tại
                            col1, col2, col3 = st.columns([2, 1, 1])
                            with col1:
                                st.text(f"✅ {filename}")
                            with col2:
                                with open(file_path, 'rb') as f:
                                    st.download_button(
                                        "📄 DOCX",
                                        data=f.read(),
                                        file_name=filename,
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key=f"download_docx_{i}"
                                    )
                        
                
                if error_list:
                    st.warning(f"⚠️ {len(error_list)} file có lỗi đã bị bỏ qua. Vui lòng sửa lỗi và thử lại.")
                
                # Cleanup processed files after download
                try:
                    cleanup_session_files()
                except:
                    pass
            else:
                st.markdown("""
                <div class="error-box">
                    <h3>❌ Xử Lý Thất Bại</h3>
                    <p>Không có file nào được xử lý thành công!</p>
                </div>
                """, unsafe_allow_html=True)
    
    elif data_list and not template_path:
        st.markdown("""
        <div class="info-box">
            <h3>⏳ Chờ Template</h3>
            <p>Vui lòng upload file template để tiếp tục</p>
        </div>
        """, unsafe_allow_html=True)
    elif not data_list and template_path:
        if error_list:
            st.markdown("""
            <div class="error-box">
                <h3>❌ Tất Cả File Đều Có Lỗi</h3>
                <p>Vui lòng kiểm tra và upload lại file hợp lệ</p>
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div class="info-box">
                <h3>⏳ Chờ File Dữ Liệu</h3>
                <p>Vui lòng upload các file dữ liệu</p>
            </div>
            """, unsafe_allow_html=True)
    else:
        st.markdown("""
        <div class="info-box">
            <h3>🚀 Bắt Đầu</h3>
            <p>Vui lòng upload file dữ liệu và template để bắt đầu</p>
        </div>
        """, unsafe_allow_html=True)
    
    # Render footer
    render_footer()

if __name__ == "__main__":
    main()
